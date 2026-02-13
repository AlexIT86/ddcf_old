# convert_manual.py
import os
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def markdown_to_docx(md_text, output_file):
    # Parsare markdown
    html = markdown.markdown(md_text, extensions=['markdown.extensions.tables'])

    # Creare document Word
    doc = Document()

    # Setare stil document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Parsing și conversie
    sections = md_text.split('##')

    # Procesare titlu principal (dacă există)
    title_match = re.search(r'^# (.+?)$', md_text, re.MULTILINE)
    if title_match:
        title = title_match.group(1)
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Procesare secțiuni
    for section in sections:
        if not section.strip():
            continue

        # Extrage titlul secțiunii (prima linie)
        lines = section.strip().split('\n')
        section_title = lines[0].strip()
        section_content = '\n'.join(lines[1:]).strip()

        # Adaugă titlul secțiunii
        doc.add_heading(section_title, 2)

        # Procesare subsecțiuni
        subsections = section_content.split('###')

        for i, subsection in enumerate(subsections):
            if not subsection.strip():
                continue

            if i == 0:  # Prima parte nu este o subsecțiune, ci conținutul secțiunii
                paragraphs = subsection.strip().split('\n\n')
                for p in paragraphs:
                    if p.strip():
                        doc.add_paragraph(p.strip())
            else:
                # Extrage titlul și conținutul subsecțiunii
                sublines = subsection.strip().split('\n')
                subsection_title = sublines[0].strip()
                subsection_content = '\n'.join(sublines[1:]).strip()

                # Adaugă titlul subsecțiunii
                doc.add_heading(subsection_title, 3)

                # Adaugă paragrafele
                paragraphs = subsection_content.split('\n\n')
                for p in paragraphs:
                    if p.strip():
                        # Verifică dacă este o listă
                        if p.strip().startswith('- '):
                            items = p.strip().split('\n- ')
                            for item in items:
                                if item.strip():
                                    para = doc.add_paragraph(style='List Bullet')
                                    para.add_run(item.strip().lstrip('- '))
                        else:
                            doc.add_paragraph(p.strip())

    # Salvare document
    doc.save(output_file)
    print(f"Documentul a fost salvat ca {output_file}")


# Citește fișierul markdown
with open('manual_utilizare.md', 'r', encoding='utf-8') as file:
    md_content = file.read()

# Convertește și salvează ca DOCX
markdown_to_docx(md_content, 'manual_utilizare.docx')